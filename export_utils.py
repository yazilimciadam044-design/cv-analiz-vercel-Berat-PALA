"""
Dışa Aktarma Modülü — PDF, Word ve CSV export fonksiyonları
"""

import json
import io
from datetime import datetime


# ─── PDF ──────────────────────────────────────────────────────────────────────

def generate_pdf(isim: str, meslek_alani: str, analiz: dict) -> bytes:
    """Analiz sonucunu PDF olarak döner (bytes)."""
    try:
        from fpdf import FPDF
    except ImportError:
        raise ImportError("fpdf2 paketi yüklü değil.")

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    # Türkçe karakter desteği için core font kullan
    pdf.set_font("Helvetica", "B", 20)
    pdf.set_fill_color(30, 30, 60)
    pdf.set_text_color(255, 255, 255)
    pdf.cell(0, 14, "  AI CV Analiz Raporu", fill=True, ln=True)
    pdf.ln(4)

    pdf.set_text_color(30, 30, 60)
    pdf.set_font("Helvetica", "", 11)
    pdf.cell(0, 8, f"Ad Soyad : {_safe(isim)}", ln=True)
    pdf.cell(0, 8, f"Meslek Alani: {_safe(meslek_alani)}", ln=True)
    pdf.cell(0, 8, f"Tarih     : {datetime.now().strftime('%d.%m.%Y %H:%M')}", ln=True)
    pdf.ln(4)

    # Puan
    puan = analiz.get("puan", 0)
    pdf.set_font("Helvetica", "B", 14)
    _section_header(pdf, f"CV Puani: {puan}/100")
    _puan_bar(pdf, puan)
    pdf.ln(3)

    # Güçlü Yönler
    _section_header(pdf, "Guclu Yonler")
    pdf.set_font("Helvetica", "", 11)
    for item in analiz.get("guclu_yonler", []):
        pdf.cell(8)
        pdf.cell(0, 7, f"+ {_safe(item)}", ln=True)
    pdf.ln(2)

    # Eksikler
    _section_header(pdf, "Eksikler")
    pdf.set_font("Helvetica", "", 11)
    for item in analiz.get("eksikler", []):
        pdf.cell(8)
        pdf.cell(0, 7, f"- {_safe(item)}", ln=True)
    pdf.ln(2)

    # Önerilen Meslekler
    _section_header(pdf, "Onerilen Meslekler")
    pdf.set_font("Helvetica", "", 11)
    meslekler = " | ".join(analiz.get("onerilen_meslekler", []))
    pdf.multi_cell(0, 7, _safe(meslekler))
    pdf.ln(2)

    # Öneriler
    _section_header(pdf, "Kariyer Onerileri")
    pdf.set_font("Helvetica", "", 11)
    for i, item in enumerate(analiz.get("oneriler", []), 1):
        pdf.cell(8)
        pdf.multi_cell(0, 7, f"{i}. {_safe(item)}")
    pdf.ln(4)

    pdf.set_font("Helvetica", "I", 9)
    pdf.set_text_color(120, 120, 120)
    pdf.cell(0, 6, "Bu rapor AI CV Analiz Sistemi tarafindan otomatik olarak olusturulmustur.", ln=True, align="C")

    return bytes(pdf.output())


def _safe(text: str) -> str:
    """Türkçe karakterleri ASCII'ye dönüştürür (FPDF core font uyumu için)."""
    replacements = {
        "ğ": "g", "Ğ": "G", "ü": "u", "Ü": "U", "ş": "s", "Ş": "S",
        "ı": "i", "İ": "I", "ö": "o", "Ö": "O", "ç": "c", "Ç": "C",
    }
    for tr, en in replacements.items():
        text = text.replace(tr, en)
    return text


def _section_header(pdf, title: str):
    pdf.set_font("Helvetica", "B", 12)
    pdf.set_fill_color(240, 240, 255)
    pdf.set_text_color(30, 30, 80)
    pdf.cell(0, 9, f"  {title}", fill=True, ln=True)
    pdf.set_text_color(30, 30, 60)
    pdf.ln(1)


def _puan_bar(pdf, puan: int):
    """Puan çubuğu çizer."""
    bar_w = 160
    bar_h = 8
    x = pdf.get_x() + 8
    y = pdf.get_y()
    # Arka plan
    pdf.set_fill_color(220, 220, 220)
    pdf.rect(x, y, bar_w, bar_h, "F")
    # Dolu kısım
    if puan >= 75:
        pdf.set_fill_color(46, 204, 113)
    elif puan >= 50:
        pdf.set_fill_color(243, 156, 18)
    else:
        pdf.set_fill_color(231, 76, 60)
    pdf.rect(x, y, bar_w * puan / 100, bar_h, "F")
    pdf.ln(bar_h + 2)


# ─── Word ─────────────────────────────────────────────────────────────────────

def generate_docx(isim: str, meslek_alani: str, analiz: dict) -> bytes:
    """Analiz sonucunu Word dosyası olarak döner (bytes)."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise ImportError("python-docx paketi yüklü değil.")

    doc = Document()

    # Başlık
    title = doc.add_heading("AI CV Analiz Raporu", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Meta bilgi
    doc.add_paragraph(f"Ad Soyad: {isim}")
    doc.add_paragraph(f"Meslek Alanı: {meslek_alani}")
    doc.add_paragraph(f"Tarih: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
    doc.add_paragraph()

    puan = analiz.get("puan", 0)
    p = doc.add_heading(f"CV Puanı: {puan}/100", level=1)

    doc.add_heading("Güçlü Yönler ✅", level=2)
    for item in analiz.get("guclu_yonler", []):
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Eksikler ❌", level=2)
    for item in analiz.get("eksikler", []):
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Önerilen Meslekler 🎯", level=2)
    doc.add_paragraph(" | ".join(analiz.get("onerilen_meslekler", [])))

    doc.add_heading("Kariyer Önerileri 💡", level=2)
    for item in analiz.get("oneriler", []):
        doc.add_paragraph(item, style="List Number")

    doc.add_paragraph()
    footer_p = doc.add_paragraph("Bu rapor AI CV Analiz Sistemi tarafından otomatik olarak oluşturulmuştur.")
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─── CSV ──────────────────────────────────────────────────────────────────────

def generate_csv(rows: list[dict]) -> bytes:
    """Analiz geçmişini CSV olarak döner."""
    import csv

    records = []
    for row in rows:
        analiz = row.get("analiz_sonucu", {})
        if not isinstance(analiz, dict):
            analiz = {}
        records.append({
            "ID": row.get("id", ""),
            "İsim": row.get("isim", ""),
            "Tarih": row.get("tarih", ""),
            "Meslek Alanı": row.get("meslek_alani", ""),
            "AI Modeli": row.get("ai_model", ""),
            "Puan": row.get("puan", ""),
            "Güçlü Yönler": "; ".join(analiz.get("guclu_yonler", [])),
            "Eksikler": "; ".join(analiz.get("eksikler", [])),
            "Önerilen Meslekler": "; ".join(analiz.get("onerilen_meslekler", [])),
            "Öneriler": "; ".join(analiz.get("oneriler", [])),
        })

    buf = io.StringIO()
    if records:
        fieldnames = list(records[0].keys())
        writer = csv.DictWriter(buf, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows(records)

    # Encode with utf-8-sig for Excel compatibility
    return buf.getvalue().encode("utf-8-sig")
